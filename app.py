import streamlit as st
import openai
import anthropic
from io import BytesIO
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import re
from collections import Counter

# Page config
st.set_page_config(
    page_title="Resume ATS Enhancer",
    page_icon="📄",
    layout="wide"
)

# Initialize session state
if 'original_resume' not in st.session_state:
    st.session_state.original_resume = ""
if 'enhanced_resume' not in st.session_state:
    st.session_state.enhanced_resume = ""
if 'job_description' not in st.session_state:
    st.session_state.job_description = ""
if 'original_ats_score' not in st.session_state:
    st.session_state.original_ats_score = 0
if 'enhanced_ats_score' not in st.session_state:
    st.session_state.enhanced_ats_score = 0


# Helper Functions
def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""


def calculate_ats_score(resume_text, job_description):
    """Calculate ATS score based on multiple factors"""
    if not resume_text or not job_description:
        return 0

    score = 0
    max_score = 100

    # Normalize text
    resume_lower = resume_text.lower()
    jd_lower = job_description.lower()

    # 1. Keyword Matching (40 points)
    # Extract important words from job description (excluding common words)
    common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                    'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'be', 'been',
                    'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
                    'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those'}

    jd_words = [word.strip('.,!?;:()[]{}') for word in jd_lower.split()]
    jd_words = [word for word in jd_words if len(word) > 2 and word not in common_words]
    jd_word_freq = Counter(jd_words)

    # Get top keywords from job description
    top_keywords = [word for word, _ in jd_word_freq.most_common(30)]

    matched_keywords = sum(1 for keyword in top_keywords if keyword in resume_lower)
    keyword_score = min(40, (matched_keywords / len(top_keywords)) * 40) if top_keywords else 0
    score += keyword_score

    # 2. Essential Sections (25 points)
    sections = {
        'experience': r'(experience|work history|employment|professional experience)',
        'education': r'(education|academic|degree|university|college)',
        'skills': r'(skills|technical skills|competencies|proficiencies)',
        'contact': r'(email|phone|linkedin|contact)'
    }

    section_score = 0
    for section, pattern in sections.items():
        if re.search(pattern, resume_lower):
            section_score += 6.25
    score += section_score

    # 3. Action Verbs (15 points)
    action_verbs = ['achieved', 'improved', 'developed', 'managed', 'led', 'created',
                    'implemented', 'designed', 'built', 'increased', 'decreased', 'launched',
                    'delivered', 'optimized', 'streamlined', 'coordinated', 'executed',
                    'spearheaded', 'founded', 'established', 'drove', 'generated']

    action_verb_count = sum(1 for verb in action_verbs if verb in resume_lower)
    action_verb_score = min(15, (action_verb_count / 10) * 15)
    score += action_verb_score

    # 4. Quantifiable Achievements (10 points)
    # Look for numbers, percentages, dollar amounts
    numbers_pattern = r'\d+%|\$\d+|(\d+\+|\d+ years|\d+ months)'
    quantifiable_count = len(re.findall(numbers_pattern, resume_text))
    quantifiable_score = min(10, (quantifiable_count / 5) * 10)
    score += quantifiable_score

    # 5. Resume Length (5 points) - optimal is 400-800 words
    word_count = len(resume_text.split())
    if 400 <= word_count <= 800:
        length_score = 5
    elif 300 <= word_count < 400 or 800 < word_count <= 1000:
        length_score = 3
    else:
        length_score = 1
    score += length_score

    # 6. Professional Formatting (5 points)
    format_score = 0
    if re.search(r'[A-Z][a-z]+\s+[A-Z][a-z]+', resume_text):  # Has proper name format
        format_score += 2
    if '@' in resume_text:  # Has email
        format_score += 1.5
    if re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', resume_text):  # Has phone
        format_score += 1.5
    score += format_score

    return round(min(score, max_score), 1)


def call_openai(resume, job_desc, model="gpt-4"):
    """Call OpenAI API to enhance resume"""
    try:
        client = openai.OpenAI(api_key=st.session_state.openai_key)

        prompt = f"""You are an expert resume writer and ATS optimization specialist. 

Job Description:
{job_desc}

Current Resume:
{resume}

Task: Enhance this resume to better match the job description and improve ATS score. Follow these guidelines:
1. Keep all factual information accurate - DO NOT fabricate experience or skills
2. Incorporate relevant keywords from the job description naturally
3. Restructure bullet points to highlight relevant achievements
4. Use action verbs and quantifiable results where possible
5. Optimize formatting for ATS systems
6. Maintain the original tone and professionalism
7. Keep the same overall structure but improve content

Return ONLY the enhanced resume text, no additional commentary."""

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert resume optimization assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )

        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error calling OpenAI: {str(e)}"


def call_claude(resume, job_desc, model="claude-sonnet-4-20250514"):
    """Call Claude API to enhance resume"""
    try:
        client = anthropic.Anthropic(api_key=st.session_state.claude_key)

        prompt = f"""You are an expert resume writer and ATS optimization specialist. 

Job Description:
{job_desc}

Current Resume:
{resume}

Task: Enhance this resume to better match the job description and improve ATS score. Follow these guidelines:
1. Keep all factual information accurate - DO NOT fabricate experience or skills
2. Incorporate relevant keywords from the job description naturally
3. Restructure bullet points to highlight relevant achievements
4. Use action verbs and quantifiable results where possible
5. Optimize formatting for ATS systems
6. Maintain the original tone and professionalism
7. Keep the same overall structure but improve content

Return ONLY the enhanced resume text, no additional commentary."""

        response = client.messages.create(
            model=model,
            max_tokens=2000,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        return response.content[0].text.strip()
    except Exception as e:
        return f"Error calling Claude: {str(e)}"


def create_pdf(text):
    """Create PDF from text"""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)

    styles = getSampleStyleSheet()
    style = ParagraphStyle(
        'CustomStyle',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )

    story = []

    # Split text into paragraphs
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            # Clean text for PDF
            para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(para, style))
        else:
            story.append(Spacer(1, 0.2 * inch))

    doc.build(story)
    buffer.seek(0)
    return buffer


def create_word(text):
    """Create Word document from text"""
    doc = Document()

    # Add paragraphs
    for line in text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            # Set font
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
        else:
            doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Main App
st.title("📄 Resume ATS Enhancer")
st.markdown("---")

# Top Bar - Model Selection and API Keys
col1, col2, col3 = st.columns([2, 2, 2])

with col1:
    llm_provider = st.selectbox(
        "🤖 Select LLM Provider",
        ["OpenAI (ChatGPT)", "Anthropic (Claude)"],
        key="llm_provider"
    )

with col2:
    if "OpenAI" in llm_provider:
        model_choice = st.selectbox(
            "Select Model",
            ["gpt-4", "gpt-4-turbo", "gpt-3.5-turbo"],
            key="model"
        )
    else:
        model_choice = st.selectbox(
            "Select Model",
            ["claude-sonnet-4-20250514", "claude-opus-4-20250514", "claude-3-5-sonnet-20241022"],
            key="model"
        )

with col3:
    if "OpenAI" in llm_provider:
        openai_key = st.text_input("OpenAI API Key", type="password", key="openai_key_input")
        if openai_key:
            st.session_state.openai_key = openai_key
    else:
        claude_key = st.text_input("Claude API Key", type="password", key="claude_key_input")
        if claude_key:
            st.session_state.claude_key = claude_key

st.markdown("---")

# ATS Score Display
if st.session_state.original_resume or st.session_state.enhanced_resume:
    st.subheader("🎯 ATS Score Comparison")

    score_col1, score_col2, score_col3 = st.columns([1, 1, 1])

    with score_col1:
        original_score = st.session_state.original_ats_score
        score_color = "#ff4444" if original_score < 50 else "#ffa500" if original_score < 70 else "#00cc00"
        st.markdown(f"""
        <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid {score_color};'>
            <h3 style='margin: 0; color: #333;'>Original Resume</h3>
            <h1 style='margin: 10px 0; color: {score_color}; font-size: 48px;'>{original_score}</h1>
            <p style='margin: 0; color: #666;'>ATS Score</p>
        </div>
        """, unsafe_allow_html=True)

    with score_col2:
        if st.session_state.enhanced_resume:
            enhanced_score = st.session_state.enhanced_ats_score
            score_color = "#ff4444" if enhanced_score < 50 else "#ffa500" if enhanced_score < 70 else "#00cc00"

            # Calculate improvement
            improvement = enhanced_score - original_score
            improvement_text = f"+{improvement}" if improvement > 0 else f"{improvement}"
            improvement_color = "#00cc00" if improvement > 0 else "#ff4444" if improvement < 0 else "#666"

            st.markdown(f"""
            <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid {score_color};'>
                <h3 style='margin: 0; color: #333;'>Enhanced Resume</h3>
                <h1 style='margin: 10px 0; color: {score_color}; font-size: 48px;'>{enhanced_score}</h1>
                <p style='margin: 0; color: #666;'>ATS Score</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid #cccccc;'>
                <h3 style='margin: 0; color: #333;'>Enhanced Resume</h3>
                <h1 style='margin: 10px 0; color: #999; font-size: 48px;'>--</h1>
                <p style='margin: 0; color: #666;'>Pending</p>
            </div>
            """, unsafe_allow_html=True)

    with score_col3:
        if st.session_state.enhanced_resume:
            improvement = st.session_state.enhanced_ats_score - st.session_state.original_ats_score
            improvement_text = f"+{improvement}" if improvement > 0 else f"{improvement}"
            improvement_color = "#00cc00" if improvement > 0 else "#ff4444" if improvement < 0 else "#666"
            arrow = "↑" if improvement > 0 else "↓" if improvement < 0 else "→"

            st.markdown(f"""
            <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid {improvement_color};'>
                <h3 style='margin: 0; color: #333;'>Improvement</h3>
                <h1 style='margin: 10px 0; color: {improvement_color}; font-size: 48px;'>{arrow} {improvement_text}</h1>
                <p style='margin: 0; color: #666;'>Points</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid #cccccc;'>
                <h3 style='margin: 0; color: #333;'>Improvement</h3>
                <h1 style='margin: 10px 0; color: #999; font-size: 48px;'>--</h1>
                <p style='margin: 0; color: #666;'>Pending</p>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("---")

# Upload Section
st.subheader("📤 Upload Documents")
col_upload1, col_upload2 = st.columns(2)

with col_upload1:
    st.write("**Upload Your Resume**")
    resume_option = st.radio("Resume Format:", ["PDF", "Text"], horizontal=True)

    if resume_option == "PDF":
        resume_file = st.file_uploader("Choose PDF file", type=['pdf'], key="resume_pdf")
        if resume_file:
            st.session_state.original_resume = extract_text_from_pdf(resume_file)
            if st.session_state.job_description:
                st.session_state.original_ats_score = calculate_ats_score(
                    st.session_state.original_resume,
                    st.session_state.job_description
                )
            st.success("✅ Resume uploaded successfully!")
    else:
        resume_text = st.text_area("Paste your resume text:", height=150, key="resume_text")
        if resume_text:
            st.session_state.original_resume = resume_text
            if st.session_state.job_description:
                st.session_state.original_ats_score = calculate_ats_score(
                    st.session_state.original_resume,
                    st.session_state.job_description
                )

with col_upload2:
    st.write("**Job Description**")
    job_desc = st.text_area("Paste the job description:", height=200, key="job_desc")
    if job_desc:
        st.session_state.job_description = job_desc
        # Recalculate scores if resume already exists
        if st.session_state.original_resume:
            st.session_state.original_ats_score = calculate_ats_score(
                st.session_state.original_resume,
                st.session_state.job_description
            )
        if st.session_state.enhanced_resume:
            st.session_state.enhanced_ats_score = calculate_ats_score(
                st.session_state.enhanced_resume,
                st.session_state.job_description
            )

# Enhance Button
st.markdown("---")
col_btn1, col_btn2, col_btn3 = st.columns([2, 2, 2])

with col_btn1:
    if st.button("🚀 Enhance Resume", type="primary", use_container_width=True):
        if not st.session_state.original_resume:
            st.error("❌ Please upload your resume first!")
        elif not st.session_state.job_description:
            st.error("❌ Please provide a job description!")
        elif "OpenAI" in llm_provider and not st.session_state.get('openai_key'):
            st.error("❌ Please provide your OpenAI API key!")
        elif "Claude" in llm_provider and not st.session_state.get('claude_key'):
            st.error("❌ Please provide your Claude API key!")
        else:
            with st.spinner("✨ Enhancing your resume..."):
                if "OpenAI" in llm_provider:
                    result = call_openai(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )
                else:
                    result = call_claude(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )

                st.session_state.enhanced_resume = result
                # Calculate ATS score for enhanced resume
                st.session_state.enhanced_ats_score = calculate_ats_score(
                    st.session_state.enhanced_resume,
                    st.session_state.job_description
                )
                st.success("✅ Resume enhanced successfully!")
                st.rerun()

with col_btn2:
    if st.button("🔄 Retry Enhancement", use_container_width=True):
        if st.session_state.enhanced_resume:
            with st.spinner("✨ Re-enhancing your resume..."):
                if "OpenAI" in llm_provider:
                    result = call_openai(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )
                else:
                    result = call_claude(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )

                st.session_state.enhanced_resume = result
                # Recalculate ATS score
                st.session_state.enhanced_ats_score = calculate_ats_score(
                    st.session_state.enhanced_resume,
                    st.session_state.job_description
                )
                st.success("✅ Resume re-enhanced successfully!")
                st.rerun()
        else:
            st.warning("⚠️ Please enhance the resume first!")

with col_btn3:
    st.write("")  # Spacing

# Display Section
st.markdown("---")
st.subheader("📊 Compare Resumes")

# Create three columns
left_col, right_col = st.columns([1, 1])

with left_col:
    # Job Description (Top)
    st.markdown("### 📋 Job Description")
    with st.container():
        st.text_area(
            "Job Description",
            value=st.session_state.job_description,
            height=250,
            disabled=True,
            key="jd_display",
            label_visibility="collapsed"
        )

    st.markdown("---")

    # Original Resume (Bottom)
    st.markdown("### 📄 Original Resume")
    with st.container():
        st.text_area(
            "Original Resume",
            value=st.session_state.original_resume,
            height=300,
            disabled=True,
            key="original_display",
            label_visibility="collapsed"
        )

with right_col:
    st.markdown("### ✨ Enhanced Resume")
    with st.container():
        enhanced_text = st.text_area(
            "Enhanced Resume",
            value=st.session_state.enhanced_resume,
            height=580,
            key="enhanced_display",
            label_visibility="collapsed"
        )

        # Update session state if user edits
        if enhanced_text != st.session_state.enhanced_resume:
            st.session_state.enhanced_resume = enhanced_text

# Download Section
if st.session_state.enhanced_resume:
    st.markdown("---")
    st.subheader("💾 Download Enhanced Resume")

    col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 2])

    with col_dl1:
        pdf_buffer = create_pdf(st.session_state.enhanced_resume)
        st.download_button(
            label="📥 Download as PDF",
            data=pdf_buffer,
            file_name="enhanced_resume.pdf",
            mime="application/pdf",
            use_container_width=True
        )

    with col_dl2:
        word_buffer = create_word(st.session_state.enhanced_resume)
        st.download_button(
            label="📥 Download as Word",
            data=word_buffer,
            file_name="enhanced_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: #666; padding: 20px;'>
        <p>💡 <strong>Tip:</strong> Always review the enhanced resume to ensure accuracy before using it for applications.</p>
        <p><strong>📊 ATS Score Guide:</strong> 
            <span style='color: #00cc00;'>70-100 = Excellent</span> | 
            <span style='color: #ffa500;'>50-69 = Good</span> | 
            <span style='color: #ff4444;'>Below 50 = Needs Improvement</span>
        </p>
        <p style='font-size: 11px; margin-top: 10px;'>
            <strong>Score Factors:</strong> Keyword Match (40%) • Essential Sections (25%) • Action Verbs (15%) • Quantifiable Results (10%) • Length (5%) • Formatting (5%)
        </p>
        <p style='font-size: 12px;'>Made with ❤️ for job seekers by NARNE| Powered by AI</p>
    </div>
    """,
    unsafe_allow_html=True
)