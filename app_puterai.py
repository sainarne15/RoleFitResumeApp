import streamlit as st
import streamlit.components.v1 as components  # NEW: for the browser-side Puter.js widget
import openai
import anthropic
from io import BytesIO
import PyPDF2
from collections import Counter
import re
import json
from datetime import datetime
import difflib
from docx import Document
import requests

# Page config
st.set_page_config(
    page_title="Role Fit Resume PRO",
    page_icon="📄",
    layout="wide"
)

# ---- Curated fallback models for Puter (used if live fetch fails or no key) ----
PUTER_MODELS_FALLBACK = [
    # OpenAI family routed by Puter
    "gpt-5-nano", "gpt-5-mini", "gpt-5", "gpt-5-chat-latest",
    "gpt-4o", "gpt-4o-mini",
    "gpt-4.1", "gpt-4.1-mini", "gpt-4.1-nano",
    "gpt-4.5-preview",
    "o1", "o1-mini", "o1-pro", "o3", "o3-mini", "o4-mini",
    # Anthropic family
    "claude-3-7-sonnet", "claude-sonnet-4-5", "claude-opus-4-1", "claude-haiku-4-5",
]

@st.cache_data(ttl=3600)
def fetch_puter_models(api_key: str):
    """Try to fetch full model list from Puter; fallback to curated list."""
    try:
        if not api_key:
            return PUTER_MODELS_FALLBACK
        resp = requests.get(
            "https://api.puter.com/v1/models",
            headers={"Authorization": f"Bearer {api_key}"},
            timeout=20
        )
        if resp.status_code != 200:
            return PUTER_MODELS_FALLBACK
        data = resp.json()
        ids = [m.get("id") for m in data.get("data", []) if m.get("id")]
        if not ids:
            return PUTER_MODELS_FALLBACK
        preferred_first = [m for m in PUTER_MODELS_FALLBACK if m in ids]
        others = [m for m in ids if m not in preferred_first]
        return preferred_first + others
    except Exception:
        return PUTER_MODELS_FALLBACK

# Initialize session state
if 'original_resume' not in st.session_state:
    st.session_state.original_resume = ""
if 'enhanced_resume' not in st.session_state:
    st.session_state.enhanced_resume = ""
if 'job_description' not in st.session_state:
    st.session_state.job_description = ""
if 'original_ats_score' not in st.session_state:
    st.session_state.original_ats_score = 0
if 'enhanced_ats_score' not in st.session_state:
    st.session_state.enhanced_ats_score = 0
if 'api_keys' not in st.session_state:
    st.session_state.api_keys = {'openai': '', 'claude': '', 'puter': ''}
if 'history' not in st.session_state:
    st.session_state.history = []
if 'current_version' not in st.session_state:
    st.session_state.current_version = 0
if 'resume_filename' not in st.session_state:
    st.session_state.resume_filename = ""

# Helper Functions
def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file with improved handling"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")

        if not text.strip():
            st.error("No text could be extracted from the PDF. The PDF might be image-based or corrupted.")
            return ""

        return text.strip()
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_word(docx_file):
    """Extract text from uploaded Word document"""
    try:
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        if not text.strip():
            st.error("No text could be extracted from the Word document.")
            return ""

        return text.strip()
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return ""

def calculate_ats_score(resume_text, job_description):
    """Calculate ATS score based on multiple factors"""
    if not resume_text or not job_description:
        return 0

    score = 0
    max_score = 100

    # Normalize text
    resume_lower = resume_text.lower()
    jd_lower = job_description.lower()

    # 1. Keyword Matching (40 points)
    common_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
                    'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'be', 'been',
                    'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could',
                    'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those'}

    jd_words = [word.strip('.,!?;:()[]{}') for word in jd_lower.split()]
    jd_words = [word for word in jd_words if len(word) > 2 and word not in common_words]
    jd_word_freq = Counter(jd_words)

    top_keywords = [word for word, _ in jd_word_freq.most_common(30)]

    matched_keywords = sum(1 for keyword in top_keywords if keyword in resume_lower)
    keyword_score = min(40, (matched_keywords / len(top_keywords)) * 40) if top_keywords else 0
    score += keyword_score

    # 2. Essential Sections (25 points)
    sections = {
        'experience': r'(experience|work history|employment|professional experience)',
        'education': r'(education|academic|degree|university|college)',
        'skills': r'(skills|technical skills|competencies|proficiencies)',
        'contact': r'(email|phone|linkedin|contact)'
    }

    section_score = 0
    for section, pattern in sections.items():
        if re.search(pattern, resume_lower):
            section_score += 6.25
    score += section_score

    # 3. Action Verbs (15 points)
    action_verbs = ['achieved', 'improved', 'developed', 'managed', 'led', 'created',
                    'implemented', 'designed', 'built', 'increased', 'decreased', 'launched',
                    'delivered', 'optimized', 'streamlined', 'coordinated', 'executed',
                    'spearheaded', 'founded', 'established', 'drove', 'generated']

    action_verb_count = sum(1 for verb in action_verbs if verb in resume_lower)
    action_verb_score = min(15, (action_verb_count / 10) * 15)
    score += action_verb_score

    # 4. Quantifiable Achievements (10 points)
    numbers_pattern = r'\d+%|\$\d+|(\d+\+|\d+ years|\d+ months)'
    quantifiable_count = len(re.findall(numbers_pattern, resume_text))
    quantifiable_score = min(10, (quantifiable_count / 5) * 10)
    score += quantifiable_score

    # 5. Resume Length (5 points)
    word_count = len(resume_text.split())
    if 400 <= word_count <= 800:
        length_score = 5
    elif 300 <= word_count < 400 or 800 < word_count <= 1000:
        length_score = 3
    else:
        length_score = 1
    score += length_score

    # 6. Professional Formatting (5 points)
    format_score = 0
    if re.search(r'[A-Z][a-z]+\s+[A-Z][a-z]+', resume_text):
        format_score += 2
    if '@' in resume_text:
        format_score += 1.5
    if re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', resume_text):
        format_score += 1.5
    score += format_score

    return round(min(score, max_score), 1)

def call_puter(resume, job_desc, model="gpt-5-nano"):
    """Server-side Puter REST call (OpenAI-compatible). Requires a Puter API key."""
    try:
        api_key = st.session_state.api_keys.get('puter', '').strip()
        if not api_key:
            return "Error: Puter API key is required for server-side calls."

        original_word_count = len(resume.split())
        original_line_count = len([line for line in resume.split('\n') if line.strip()])

        prompt = f"""You are an expert resume writer and ATS optimization specialist with a focus on SURGICAL, MINIMAL changes.

Job Description:
{job_desc}

Current Resume (Word count: {original_word_count}, Lines: {original_line_count}):
{resume}

CRITICAL INSTRUCTIONS - READ CAREFULLY:
1) LENGTH: Target exactly {original_word_count}±30 words and {original_line_count}±3 lines. If you add, condense elsewhere.
2) SELECTIVE changes: Improve only weak/irrelevant bullets; leave strong ones untouched.
3) MODIFY: add metrics, weave relevant JD keywords, upgrade weak verbs, compress irrelevant parts.
4) DO NOT FABRICATE; preserve sections/titles/dates/formatting.
Return ONLY the enhanced resume with the same structure.
"""

        headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": "You are an expert resume optimizer who makes minimal, targeted changes."},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.5,
            "max_tokens": 2500,
        }

        resp = requests.post("https://api.puter.com/v1/chat/completions", headers=headers, json=payload, timeout=90)
        if resp.status_code == 200:
            data = resp.json()
            return data["choices"][0]["message"]["content"].strip()
        else:
            return f"Error calling Puter AI: {resp.status_code} - {resp.text}"

    except Exception as e:
        return f"Error calling Puter AI: {str(e)}"

def call_openai(resume, job_desc, model="gpt-4o"):
    """Direct OpenAI call."""
    try:
        client = openai.OpenAI(api_key=st.session_state.api_keys.get('openai', '').strip())

        original_word_count = len(resume.split())
        original_line_count = len([line for line in resume.split('\n') if line.strip()])

        prompt = f"""You are an expert resume writer and ATS optimization specialist with a focus on SURGICAL, MINIMAL changes.

Job Description:
{job_desc}

Current Resume (Word count: {original_word_count}, Lines: {original_line_count}):
{resume}

CRITICAL INSTRUCTIONS - READ CAREFULLY:
1) LENGTH: Target exactly {original_word_count}±30 words and {original_line_count}±3 lines. If you add, condense elsewhere.
2) SELECTIVE changes: Improve only weak/irrelevant bullets; leave strong ones untouched.
3) MODIFY: add metrics, weave relevant JD keywords, upgrade weak verbs, compress irrelevant parts.
4) DO NOT FABRICATE; preserve sections/titles/dates/formatting.
Return ONLY the enhanced resume with the same structure.
"""

        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert resume optimizer who makes minimal, targeted changes."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.5,
            max_tokens=2500,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error calling OpenAI: {str(e)}"

def call_claude(resume, job_desc, model="claude-sonnet-4-20250514"):
    """Call Claude API to enhance resume"""
    try:
        client = anthropic.Anthropic(api_key=st.session_state.api_keys['claude'])

        original_word_count = len(resume.split())
        original_line_count = len([line for line in resume.split('\n') if line.strip()])

        prompt = f"""You are an expert resume writer and ATS optimization specialist with a focus on SURGICAL, MINIMAL changes.

Job Description:
{job_desc}

Current Resume (Word count: {original_word_count}, Lines: {original_line_count}):
{resume}

CRITICAL INSTRUCTIONS - READ CAREFULLY:
1) LENGTH: Target exactly {original_word_count}±30 words and {original_line_count}±3 lines. If you add, condense elsewhere.
2) SELECTIVE changes: Improve only weak/irrelevant bullets; leave strong ones untouched.
3) MODIFY: add metrics, weave relevant JD keywords, upgrade weak verbs, compress irrelevant parts.
4) DO NOT FABRICATE; preserve sections/titles/dates/formatting.
Return ONLY the enhanced resume with the same structure.
"""

        response = client.messages.create(
            model=model,
            max_tokens=2500,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )

        return response.content[0].text.strip()
    except Exception as e:
        return f"Error calling Claude: {str(e)}"

def save_to_history(resume_text, score, version_num):
    """Save resume version to history"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.history.append({
        'version': version_num,
        'timestamp': timestamp,
        'resume': resume_text,
        'score': score,
        'words': len(resume_text.split()),
        'lines': len([l for l in resume_text.split('\n') if l.strip()])
    })

def highlight_changes(original_text, enhanced_text):
    """Highlight changes between original and enhanced text using improved diff algorithm"""
    if not original_text or not enhanced_text:
        return enhanced_text

    # Normalize whitespace for better comparison
    original_text = original_text.strip()
    enhanced_text = enhanced_text.strip()

    # Split into lines
    original_lines = original_text.split('\n')
    enhanced_lines = enhanced_text.split('\n')

    result_html = []

    # Use SequenceMatcher to compare lines
    line_matcher = difflib.SequenceMatcher(None, original_lines, enhanced_lines)

    for tag, i1, i2, j1, j2 in line_matcher.get_opcodes():
        if tag == 'equal':
            for line in enhanced_lines[j1:j2]:
                result_html.append(line)

        elif tag == 'replace':
            for orig_idx, enh_idx in zip(range(i1, i2), range(j1, j2)):
                if orig_idx < len(original_lines) and enh_idx < len(enhanced_lines):
                    orig_line = original_lines[orig_idx]
                    enh_line = enhanced_lines[enh_idx]
                    similarity = difflib.SequenceMatcher(None, orig_line, enh_line).ratio()

                    if similarity > 0.3:
                        orig_words = orig_line.split()
                        enh_words = enh_line.split()

                        word_matcher = difflib.SequenceMatcher(None, orig_words, enh_words)
                        line_parts = []

                        for word_tag, wi1, wi2, wj1, wj2 in word_matcher.get_opcodes():
                            if word_tag == 'equal':
                                line_parts.append(' '.join(enh_words[wj1:wj2]))
                            elif word_tag in ('replace', 'insert'):
                                changed_text = ' '.join(enh_words[wj1:wj2])
                                if changed_text.strip():
                                    line_parts.append(
                                        f'<span style="background-color: #90EE90; font-weight: 500; padding: 1px 3px; border-radius: 2px;">{changed_text}</span>'
                                    )
                        result_html.append(' '.join(line_parts) if line_parts else enh_line)
                    else:
                        result_html.append(
                            f'<span style="background-color: #90EE90; font-weight: 500; padding: 1px 3px; border-radius: 2px;">{enh_line}</span>'
                        )

        elif tag == 'insert':
            for line in enhanced_lines[j1:j2]:
                if line.strip():
                    result_html.append(
                        f'<span style="background-color: #90EE90; font-weight: 500; padding: 1px 3px; border-radius: 2px;">{line}</span>'
                    )
                else:
                    result_html.append(line)

        elif tag == 'delete':
            pass

    return '\n'.join(result_html)

# CSS for better styling
st.markdown("""
<style>
    .stTextArea textarea {
        font-family: 'Courier New', monospace;
        font-size: 12px;
    }
    .history-item {
        padding: 10px;
        margin: 5px 0;
        border-radius: 5px;
        border: 1px solid #ddd;
        cursor: pointer;
    }
    .history-item:hover {
        background-color: #f0f0f0;
    }
    .enhanced-display {
        font-family: 'Courier New', monospace;
        font-size: 12px;
        line-height: 1.6;
        white-space: pre-wrap;
        background-color: #f8f9fa;
        padding: 15px;
        border-radius: 5px;
        border: 1px solid #dee2e6;
        max-height: 580px;
        overflow-y: auto;
    }
</style>
""", unsafe_allow_html=True)

# Main App
st.title("📄 Resume ATS Enhancer Pro")
st.markdown("---")

# Sidebar for API Keys and Settings
with st.sidebar:
    st.header("⚙️ Settings")

    st.subheader("🔑 API Keys (Persistent)")

    # API Provider Selection
    api_provider = st.radio(
        "Select API Provider:",
        ["Puter AI (Free)", "OpenAI (Direct)", "Anthropic Claude (Direct)"],
        help="Puter AI provides free access to GPT and Claude models (server-side needs a key; browser-side widget needs no key)."
    )

    if api_provider == "Puter AI (Free)":
        # Puter API Key
        puter_key_input = st.text_input(
            "Puter AI API Key",
            value=st.session_state.api_keys['puter'],
            type="password",
            key="puter_key_sidebar",
            help="Get a free key at puter.com. Not needed for the Advanced → Puter.js (browser) widget."
        )
        if puter_key_input:
            st.session_state.api_keys['puter'] = puter_key_input
            st.success("✅ Puter key saved")

        st.info("💡 Puter AI can be used server-side (with key) or via the in-browser widget (no key).")

    elif api_provider == "OpenAI (Direct)":
        # OpenAI API Key
        openai_key_input = st.text_input(
            "OpenAI API Key",
            value=st.session_state.api_keys['openai'],
            type="password",
            key="openai_key_sidebar",
            help="Your API key will be saved for this session"
        )
        if openai_key_input:
            st.session_state.api_keys['openai'] = openai_key_input
            st.success("✅ OpenAI key saved")

    else:  # Anthropic Claude
        # Claude API Key
        claude_key_input = st.text_input(
            "Claude API Key",
            value=st.session_state.api_keys['claude'],
            type="password",
            key="claude_key_sidebar",
            help="Your API key will be saved for this session"
        )
        if claude_key_input:
            st.session_state.api_keys['claude'] = claude_key_input
            st.success("✅ Claude key saved")

    st.markdown("---")

    # History Section
    if st.session_state.history:
        st.subheader("📜 Version History")
        st.caption(f"Total versions: {len(st.session_state.history)}")

        for idx, item in enumerate(reversed(st.session_state.history)):
            with st.expander(f"V{item['version']} - Score: {item['score']} ({item['timestamp']})"):
                st.write(f"**Words:** {item['words']} | **Lines:** {item['lines']}")
                if st.button(f"Restore V{item['version']}", key=f"restore_{idx}"):
                    st.session_state.enhanced_resume = item['resume']
                    st.session_state.enhanced_ats_score = item['score']
                    st.session_state.current_version = item['version']
                    st.rerun()

    st.markdown("---")

    # Reset button
    if st.button("🔄 Reset Session", type="secondary", use_container_width=True):
        st.session_state.original_resume = ""
        st.session_state.enhanced_resume = ""
        st.session_state.job_description = ""
        st.session_state.history = []
        st.session_state.current_version = 0
        st.session_state.resume_filename = ""
        st.rerun()

# Top Bar - Model Selection
col1, col2 = st.columns([3, 3])

with col1:
    llm_provider = st.selectbox(
        "🤖 Select LLM Provider",
        ["Puter AI (Free)", "OpenAI (ChatGPT)", "Anthropic (Claude)"],
        key="llm_provider"
    )

with col2:
    if "Puter" in llm_provider:
        _models_list = fetch_puter_models(st.session_state.api_keys.get('puter', ''))
        default_idx = _models_list.index("gpt-5-nano") if "gpt-5-nano" in _models_list else 0
        model_choice = st.selectbox(
            "Select Model",
            _models_list,
            index=default_idx,
            key="model",
            help="Models routed by Puter to multiple providers. No key needed for the browser-side widget."
        )
    elif "OpenAI" in llm_provider:
        model_choice = st.selectbox(
            "Select Model",
            ["gpt-4", "gpt-4-turbo", "gpt-3.5-turbo"],
            key="model"
        )
    else:
        model_choice = st.selectbox(
            "Select Model",
            ["claude-sonnet-4-20250514", "claude-opus-4-20250514", "claude-3-5-sonnet-20241022"],
            key="model"
        )

st.markdown("---")

# ATS Score Display
if st.session_state.original_resume or st.session_state.enhanced_resume:
    st.subheader("🎯 ATS Score Comparison")

    score_col1, score_col2, score_col3 = st.columns([1, 1, 1])

    with score_col1:
        original_score = st.session_state.original_ats_score
        score_color = "#ff4444" if original_score < 50 else "#ffa500" if original_score < 70 else "#00cc00"
        st.markdown(f"""
        <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid {score_color};'>
            <h3 style='margin: 0; color: #333;'>Original Resume</h3>
            <h1 style='margin: 10px 0; color: {score_color}; font-size: 48px;'>{original_score}</h1>
            <p style='margin: 0; color: #666;'>ATS Score</p>
        </div>
        """, unsafe_allow_html=True)

    with score_col2:
        if st.session_state.enhanced_resume:
            enhanced_score = st.session_state.enhanced_ats_score
            score_color = "#ff4444" if enhanced_score < 50 else "#ffa500" if enhanced_score < 70 else "#00cc00"

            st.markdown(f"""
            <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid {score_color};'>
                <h3 style='margin: 0; color: #333;'>Enhanced Resume (V{st.session_state.current_version})</h3>
                <h1 style='margin: 10px 0; color: {score_color}; font-size: 48px;'>{enhanced_score}</h1>
                <p style='margin: 0; color: #666;'>ATS Score</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid #cccccc;'>
                <h3 style='margin: 0; color: #333;'>Enhanced Resume</h3>
                <h1 style='margin: 10px 0; color: #999; font-size: 48px;'>--</h1>
                <p style='margin: 0; color: #666;'>Pending</p>
            </div>
            """, unsafe_allow_html=True)

    with score_col3:
        if st.session_state.enhanced_resume:
            improvement = st.session_state.enhanced_ats_score - st.session_state.original_ats_score
            improvement_text = f"+{improvement}" if improvement > 0 else f"{improvement}"
            improvement_color = "#00cc00" if improvement > 0 else "#ff4444" if improvement < 0 else "#666"
            arrow = "↑" if improvement > 0 else "↓" if improvement < 0 else "→"

            st.markdown(f"""
            <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid {improvement_color};'>
                <h3 style='margin: 0; color: #333;'>Improvement</h3>
                <h1 style='margin: 10px 0; color: {improvement_color}; font-size: 48px;'>{arrow} {improvement_text}</h1>
                <p style='margin: 0; color: #666;'>Points</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style='text-align: center; padding: 20px; background-color: #f0f2f6; border-radius: 10px; border: 2px solid #cccccc;'>
                <h3 style='margin: 0; color: #333;'>Improvement</h3>
                <h1 style='margin: 10px 0; color: #999; font-size: 48px;'>--</h1>
                <p style='margin: 0; color: #666;'>Pending</p>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("---")

# Upload Section
st.subheader("📤 Upload Documents")
col_upload1, col_upload2 = st.columns(2)

with col_upload1:
    st.write("**Upload Your Resume**")
    resume_option = st.radio("Resume Format:", ["PDF", "Word", "Text"], horizontal=True)

    if resume_option == "PDF":
        resume_file = st.file_uploader("Choose PDF file", type=['pdf'], key="resume_pdf")
        if resume_file and resume_file.name != st.session_state.resume_filename:
            st.session_state.original_resume = extract_text_from_pdf(resume_file)
            if st.session_state.original_resume:
                st.session_state.resume_filename = resume_file.name
                st.session_state.history = []
                st.session_state.current_version = 0
                st.session_state.enhanced_resume = ""
                if st.session_state.job_description:
                    st.session_state.original_ats_score = calculate_ats_score(
                        st.session_state.original_resume,
                        st.session_state.job_description
                    )
                st.success("✅ PDF uploaded successfully! History reset.")
                st.rerun()

    elif resume_option == "Word":
        resume_file = st.file_uploader("Choose Word file", type=['docx', 'doc'], key="resume_word")
        if resume_file and resume_file.name != st.session_state.resume_filename:
            st.session_state.original_resume = extract_text_from_word(resume_file)
            if st.session_state.original_resume:
                st.session_state.resume_filename = resume_file.name
                st.session_state.history = []
                st.session_state.current_version = 0
                st.session_state.enhanced_resume = ""
                if st.session_state.job_description:
                    st.session_state.original_ats_score = calculate_ats_score(
                        st.session_state.original_resume,
                        st.session_state.job_description
                    )
                st.success("✅ Word document uploaded successfully! History reset.")
                st.rerun()

    else:  # Text
        resume_text = st.text_area("Paste your resume text:", height=150, key="resume_text")
        if resume_text and resume_text != st.session_state.original_resume:
            st.session_state.original_resume = resume_text
            st.session_state.history = []
            st.session_state.current_version = 0
            st.session_state.enhanced_resume = ""
            st.session_state.resume_filename = "text_input"
            if st.session_state.job_description:
                st.session_state.original_ats_score = calculate_ats_score(
                    st.session_state.original_resume,
                    st.session_state.job_description
                )

with col_upload2:
    st.write("**Job Description**")
    job_desc = st.text_area("Paste the job description:", height=200, key="job_desc")
    if job_desc:
        st.session_state.job_description = job_desc
        if st.session_state.original_resume:
            st.session_state.original_ats_score = calculate_ats_score(
                st.session_state.original_resume,
                st.session_state.job_description
            )
        if st.session_state.enhanced_resume:
            st.session_state.enhanced_ats_score = calculate_ats_score(
                st.session_state.enhanced_resume,
                st.session_state.job_description
            )

# Enhance Button
st.markdown("---")
col_btn1, col_btn2, col_btn3 = st.columns([2, 2, 2])

with col_btn1:
    if st.button("🚀 Enhance Resume", type="primary", use_container_width=True):
        if not st.session_state.original_resume:
            st.error("❌ Please upload your resume first!")
        elif not st.session_state.job_description:
            st.error("❌ Please provide a job description!")
        elif "Puter" in llm_provider and not st.session_state.api_keys['puter']:
            st.error("❌ Please provide your Puter AI API key in the sidebar for server-side calls (or use the Advanced → Puter.js widget below).")
        elif "OpenAI" in llm_provider and not st.session_state.api_keys['openai']:
            st.error("❌ Please provide your OpenAI API key in the sidebar!")
        elif "Claude" in llm_provider and not st.session_state.api_keys['claude']:
            st.error("❌ Please provide your Claude API key in the sidebar!")
        else:
            with st.spinner("✨ Enhancing your resume with minimal targeted changes..."):
                if "Puter" in llm_provider:
                    result = call_puter(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )
                elif "OpenAI" in llm_provider:
                    result = call_openai(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )
                else:
                    result = call_claude(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )

                st.session_state.enhanced_resume = result
                st.session_state.enhanced_ats_score = calculate_ats_score(
                    st.session_state.enhanced_resume,
                    st.session_state.job_description
                )
                st.session_state.current_version += 1
                save_to_history(result, st.session_state.enhanced_ats_score, st.session_state.current_version)
                st.success("✅ Resume enhanced successfully!")
                st.rerun()

with col_btn2:
    if st.button("🔄 Retry Enhancement", use_container_width=True):
        if not st.session_state.original_resume or not st.session_state.job_description:
            st.warning("⚠️ Please upload resume and job description first!")
        elif "Puter" in llm_provider and not st.session_state.api_keys['puter']:
            st.error("❌ Please provide your Puter AI API key in the sidebar for server-side calls (or use the Advanced → Puter.js widget below).")
        elif "OpenAI" in llm_provider and not st.session_state.api_keys['openai']:
            st.error("❌ Please provide your OpenAI API key in the sidebar!")
        elif "Claude" in llm_provider and not st.session_state.api_keys['claude']:
            st.error("❌ Please provide your Claude API key in the sidebar!")
        else:
            with st.spinner("✨ Generating new version..."):
                if "Puter" in llm_provider:
                    result = call_puter(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )
                elif "OpenAI" in llm_provider:
                    result = call_openai(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )
                else:
                    result = call_claude(
                        st.session_state.original_resume,
                        st.session_state.job_description,
                        model_choice
                    )

                st.session_state.enhanced_resume = result
                st.session_state.enhanced_ats_score = calculate_ats_score(
                    st.session_state.enhanced_resume,
                    st.session_state.job_description
                )
                st.session_state.current_version += 1
                save_to_history(result, st.session_state.enhanced_ats_score, st.session_state.current_version)
                st.success("✅ New version generated!")
                st.rerun()

with col_btn3:
    st.write("")  # Spacing

# Display Section
st.markdown("---")
st.subheader("📊 Compare Resumes")

# Length comparison info
if st.session_state.original_resume:
    orig_words = len(st.session_state.original_resume.split())
    orig_lines = len([l for l in st.session_state.original_resume.split('\n') if l.strip()])

    if st.session_state.enhanced_resume:
        enh_words = len(st.session_state.enhanced_resume.split())
        enh_lines = len([l for l in st.session_state.enhanced_resume.split('\n') if l.strip()])

        col_info1, col_info2, col_info3 = st.columns([1, 1, 1])
        with col_info1:
            st.info(f"📏 **Original:** {orig_words} words, {orig_lines} lines")
        with col_info2:
            st.info(f"📏 **Enhanced V{st.session_state.current_version}:** {enh_words} words, {enh_lines} lines")
        with col_info3:
            word_diff = enh_words - orig_words
            line_diff = enh_lines - orig_lines
            diff_color = "🟢" if abs(word_diff) <= 30 else "🟡" if abs(word_diff) <= 50 else "🔴"
            st.info(f"{diff_color} **Difference:** {word_diff:+d} words, {line_diff:+d} lines")
    else:
        st.info(f"📏 **Original Resume:** {orig_words} words, {orig_lines} lines")

# Create two columns
left_col, right_col = st.columns([1, 1])

with left_col:
    st.markdown("### 📋 Job Description")
    with st.container():
        st.text_area(
            "Job Description",
            value=st.session_state.job_description,
            height=250,
            disabled=True,
            key="jd_display",
            label_visibility="collapsed"
        )

    st.markdown("---")

    st.markdown("### 📄 Original Resume")
    with st.container():
        st.text_area(
            "Original Resume",
            value=st.session_state.original_resume,
            height=300,
            disabled=True,
            key="original_display",
            label_visibility="collapsed"
        )

with right_col:
    st.markdown(f"### ✨ Enhanced Resume (Version {st.session_state.current_version})")

    # Toggle for showing highlights
    col_toggle1, col_toggle2 = st.columns([3, 1])
    with col_toggle1:
        show_highlights = st.checkbox("🎨 Show Changes Highlighted", value=True, key="show_highlights")
    with col_toggle2:
        st.write("")

    with st.container():
        if st.session_state.enhanced_resume:
            if show_highlights:
                highlighted_text = highlight_changes(
                    st.session_state.original_resume,
                    st.session_state.enhanced_resume
                )
                st.markdown(
                    f'<div class="enhanced-display">{highlighted_text}</div>',
                    unsafe_allow_html=True
                )
                st.caption("🟢 Green = Modified or Added Content")
            else:
                enhanced_text = st.text_area(
                    "Enhanced Resume",
                    value=st.session_state.enhanced_resume,
                    height=580,
                    key="enhanced_display",
                    label_visibility="collapsed"
                )

                if enhanced_text != st.session_state.enhanced_resume:
                    st.session_state.enhanced_resume = enhanced_text
        else:
            st.text_area(
                "Enhanced Resume",
                value="",
                height=580,
                key="enhanced_display_empty",
                label_visibility="collapsed",
                placeholder="Enhanced resume will appear here..."
            )

    # ---- Advanced: Browser-side Puter.js (no API key) ----
    with st.expander("⚡ Advanced: Run Puter.js in your browser (no API key)"):
        st.caption("This uses Puter.js directly in your browser. Click Run, then Copy and paste the output below, and press “Use this result”.")
        _owc = len(st.session_state.original_resume.split()) if st.session_state.original_resume else 0
        _olc = len([l for l in st.session_state.original_resume.split('\n') if l.strip()]) if st.session_state.original_resume else 0
        _prompt = f"""You are an expert resume writer and ATS optimization specialist with a focus on SURGICAL, MINIMAL changes.

Job Description:
{st.session_state.job_description}

Current Resume (Word count: {_owc}, Lines: {_olc}):
{st.session_state.original_resume}

CRITICAL INSTRUCTIONS - READ CAREFULLY:
1) LENGTH: Target exactly {_owc}±30 words and {_olc}±3 lines. If you add, condense elsewhere.
2) SELECTIVE changes: Improve only weak/irrelevant bullets; leave strong ones untouched.
3) MODIFY: add metrics, weave relevant JD keywords, upgrade weak verbs, compress irrelevant parts.
4) DO NOT FABRICATE; preserve sections/titles/dates/formatting.
Return ONLY the enhanced resume with the same structure.
"""

        # Build a stable list once
        _puter_models_for_frontend = fetch_puter_models(st.session_state.api_keys.get('puter', '')) or PUTER_MODELS_FALLBACK
        default_frontend_model = model_choice if "Puter" in llm_provider and model_choice in _puter_models_for_frontend else "gpt-5-nano"
        try:
            default_frontend_index = _puter_models_for_frontend.index(default_frontend_model)
        except ValueError:
            default_frontend_index = 0

        frontend_model = st.selectbox(
            "Puter.js model (browser-side)",
            _puter_models_for_frontend,
            index=default_frontend_index,
            key="puter_js_model"
        )

        components.html(f"""
        <!DOCTYPE html>
        <html>
        <head>
          <meta charset="utf-8" />
          <script src="https://js.puter.com/v2/"></script>
          <style>
            body {{ font-family: sans-serif; }}
            .btn {{ padding: 8px 12px; border: 1px solid #888; border-radius: 6px; background:#f5f5f5; cursor:pointer; margin-right:8px; }}
            .out {{ width:100%; height:260px; white-space:pre-wrap; }}
            .row {{ margin: 6px 0; }}
          </style>
        </head>
        <body>
          <div class="row">
            <button class="btn" id="signin">🔐 Sign in to Puter</button>
            <span id="status">Not signed in</span>
          </div>
          <div class="row">
            <button class="btn" onclick="runPuter()">▶ Run Puter.js (no key)</button>
            <button class="btn" onclick="copyOut()">📋 Copy result</button>
          </div>
          <p style="margin:8px 0 4px 0;"><b>Output</b></p>
          <textarea id="out" class="out" placeholder="Result will appear here..."></textarea>

          <script>
          const model  = {json.dumps(frontend_model)};
          const prompt = {json.dumps(_prompt)};

          // Update sign-in status label
          async function updateStatus() {{
            const signed = await puter.auth.isSignedIn();
            document.getElementById('status').textContent = signed ? "Signed in" : "Not signed in";
          }}

          async function ensureAuth() {{
            const signed = await puter.auth.isSignedIn();
            if (!signed) {{
              // Try to create a temp/guest user to avoid full signup
              await puter.auth.signIn({{ attempt_temp_user_creation: true }});
            }}
          }}

          document.getElementById('signin').addEventListener('click', async () => {{
            try {{
              await puter.auth.signIn({{ attempt_temp_user_creation: true }});
              await updateStatus();
              alert("You're signed in. You can run the model now.");
            }} catch (e) {{
              alert("Sign-in failed: " + (e && e.message ? e.message : e));
            }}
          }});

          async function runPuter() {{
            const out = document.getElementById('out');
            out.value = "Preparing...";
            try {{
              // Ensure auth is completed from this user-initiated click
              await ensureAuth();
              await updateStatus();
              out.value = "Running Puter.js with model: " + model + "...";

              // Correct Puter.js call
              const r = await puter.ai.chat([
                {{ role: "system", content: "You are an expert resume optimizer who makes minimal, targeted changes." }},
                {{ role: "user", content: prompt }}
              ], {{ model: model, temperature: 0.5, max_tokens: 2500 }});

              let text = "";
              if (typeof r === "string")       text = r;
              else if (r?.message?.content)    text = r.message.content;
              else if (r?.text)                text = r.text;
              else                             text = JSON.stringify(r);
              out.value = text || "No content received.";
            }} catch (e) {{
              out.value = "Error: " + (e && e.message ? e.message : e);
            }}
          }}

          async function copyOut() {{
            const out = document.getElementById('out');
            try {{
              await navigator.clipboard.writeText(out.value);
              alert("Copied to clipboard!");
            }} catch (e) {{
              alert("Copy failed: " + (e && e.message ? e.message : e));
            }}
          }}

          updateStatus();
          </script>
        </body>
        </html>
        """, height=420, scrolling=True)

        pasted = st.text_area("Paste output here, then click “Use this result”:", height=200, key="puter_js_paste")
        if st.button("✅ Use this result", use_container_width=True):
            if pasted and pasted.strip():
                st.session_state.enhanced_resume = pasted.strip()
                st.session_state.enhanced_ats_score = calculate_ats_score(
                    st.session_state.enhanced_resume,
                    st.session_state.job_description
                )
                st.session_state.current_version += 1
                save_to_history(st.session_state.enhanced_resume, st.session_state.enhanced_ats_score, st.session_state.current_version)
                st.success("Imported browser-side Puter.js output into the app.")
                st.rerun()
            else:
                st.warning("Please paste the Puter.js output above first.")

# Download Section
if st.session_state.enhanced_resume:
    st.markdown("---")
    st.subheader("💾 Download Enhanced Resume")

    col_dl1, col_dl2 = st.columns([1, 1])

    with col_dl1:
        # Plain text download
        st.download_button(
            label="📥 Download as Text (.txt)",
            data=st.session_state.enhanced_resume,
            file_name=f"enhanced_resume_v{st.session_state.current_version}.txt",
            mime="text/plain",
            use_container_width=True,
            help="Download as plain text to maintain exact formatting"
        )

    with col_dl2:
        # Copy to clipboard info
        st.info(
            "💡 **Tip:** Copy the enhanced resume text and paste it into your original resume document to maintain formatting!")

# Footer
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: #666; padding: 20px;'>
        <p>💡 <strong>Smart Enhancement:</strong> The app makes MINIMAL, TARGETED changes - only enhancing weak points while preserving strong ones.</p>
        <p>📝 <strong>Formatting Tip:</strong> Copy the enhanced text and paste it into your original document to preserve exact formatting, fonts, and layout.</p>
        <p>🔄 <strong>Version History:</strong> Access previous versions from the sidebar to compare and restore any version you prefer.</p>
        <p><strong>📊 ATS Score Guide:</strong> 
            <span style='color: #00cc00;'>70-100 = Excellent</span> | 
            <span style='color: #ffa500;'>50-69 = Good</span> | 
            <span style='color: #ff4444;'>Below 50 = Needs Improvement</span>
        </p>
        <p style='font-size: 11px; margin-top: 10px;'>
            <strong>Score Factors:</strong> Keyword Match (40%) • Essential Sections (25%) • Action Verbs (15%) • Quantifiable Results (10%) • Length (5%) • Formatting (5%)
        </p>
        <p style='font-size: 12px; margin-top: 15px;'>Made with ❤️ by NARNE | Powered by AI</p>
    </div>
    """,
    unsafe_allow_html=True
)
